from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from images import download_images
from convert import docx_to_pdf

def create_word_document(num: int) -> None:
    """
    Create a word document with random images and their names
    params: num - number of pages with photos
    """
    document = Document()
    names = download_images(url="https://yandex.ru/images/", num=num)
    
    for index, name in enumerate(names):
        document.add_heading(name, 0)
        document.add_picture("image" + str(index+1) + ".jpg", width=Inches(5))
        document.add_page_break()
        
    document.save('example.docx')

if __name__ == '__main__':
    num = int(input())
    create_word_document(num)
    docx_to_pdf("example.docx", "example.pdf")
    
    